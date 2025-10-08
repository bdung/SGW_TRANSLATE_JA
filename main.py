# import whisper
# from docx import Document

# # 1️⃣ Load model Whisper
# model = whisper.load_model("large")

# # 2️⃣ Transcribe file audio tiếng Nhật
# result = model.transcribe(r"D:\Downloads\Tiengnhathay.com_1500-cau-giao-tiep-tieng-nhat-mp3\file 22.mp3", language="ja")


# # 4️⃣ Tạo file Word và lưu nội dung
# doc = Document()
# doc.add_heading("Transcription tiếng Nhật", level=1)
# doc.add_paragraph(result["text"])
# doc.save(r"D:\Downloads\Tiengnhathay.com_1500-cau-giao-tiep-tieng-nhat-mp3\file_22_3.docx")

# print("Đã lưu nội dung vào file Word thành công!")
import streamlit as st
import whisper
from docx import Document
import tempfile
import os

# Load model (nên chọn small/medium để deploy online nhanh hơn)
@st.cache_resource
def load_model():
    return whisper.load_model("base")

model = load_model()

st.title("🎤 Whisper Transcription App")
st.write("Upload file audio tiếng Nhật, app sẽ nhận diện và xuất ra file Word.")

# Upload file
uploaded_file = st.file_uploader("Chọn file mp3 hoặc wav", type=["mp3", "wav"])

if uploaded_file is not None:
    # Lưu file tạm
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as tmp:
        tmp.write(uploaded_file.read())
        tmp_path = tmp.name
    
    with st.spinner("⏳ Đang xử lý..."):
        # Nhận diện giọng nói
        result = model.transcribe(tmp_path, language="ja")
        text = result["text"]

        # Hiển thị text
        st.subheader("📜 Kết quả:")
        st.write(text)

        # Xuất file Word
        doc = Document()
        doc.add_heading("Transcription tiếng Nhật", level=1)
        doc.add_paragraph(text)
        output_path = "output.docx"
        doc.save(output_path)

    # Nút download
    with open(output_path, "rb") as f:
        st.download_button("⬇️ Tải về file Word", f, file_name="transcription.docx")

    # Xoá file tạm
    os.remove(tmp_path)

